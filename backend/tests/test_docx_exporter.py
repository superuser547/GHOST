from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import RGBColor

from app.models import (
    AppendixBlock,
    FigureBlock,
    ListBlock,
    Report,
    ReportMeta,
    SectionBlock,
    TableBlock,
    TextBlock,
    WorkType,
)
from app.services.docx import build_docx


def _build_minimal_valid_report() -> Report:
    meta = ReportMeta(
        preset="misis_v1",
        work_type=WorkType.PRACTICE,
        work_number=1,
        work_type_custom=None,
        discipline="Информатика",
        topic="Первая практическая работа",
        student_full_name="Иванов Иван Иванович",
        group="ББИ-24-3",
        semester="2",
        direction_code="38.03.05",
        direction_name="Бизнес-информатика",
        department="Кафедра информатики",
        teacher_full_name="Петров Пётр Петрович",
        submission_date=date(2025, 5, 20),
    )

    blocks = [
        SectionBlock(title="Введение"),
        TextBlock(text="Текст введения."),
    ]

    return Report(meta=meta, blocks=blocks)


def _build_sample_report_with_nested_blocks() -> Report:
    meta = ReportMeta(
        preset="misis_v1",
        work_type=WorkType.LAB,
        work_number=2,
        work_type_custom=None,
        discipline="Физика",
        topic="Лабораторная работа №2",
        student_full_name="Петров Петр Петрович",
        group="ББИ-24-3",
        semester="2",
        direction_code="38.03.05",
        direction_name="Бизнес-информатика",
        department="Кафедра физики",
        teacher_full_name="Иванова Анна Сергеевна",
        submission_date=date(2025, 2, 1),
    )

    intro = SectionBlock(title="Введение", special_kind="INTRO")
    intro.children.append(TextBlock(text="Краткое введение."))

    main_section = SectionBlock(title="1 Постановка задачи")
    main_section.children.append(
        TableBlock(
            caption="Таблица 1 – Пример данных",
            rows=[["Колонка 1", "Колонка 2"], ["1", "2"]],
        )
    )
    main_section.children.append(
        FigureBlock(
            caption="Рисунок 1 – Схема установки",
            file_name="figure1.png",
        )
    )
    main_section.children.append(TextBlock(text="Комментарий к рисунку и таблице."))

    appendix = AppendixBlock(
        label="А",
        title="Дополнительные материалы",
        children=[TextBlock(text="Текст приложения.")],
    )

    return Report(meta=meta, blocks=[intro, main_section, appendix])


def _build_sample_report_with_list_and_table() -> Report:
    meta = ReportMeta(
        preset="misis_v1",
        work_type=WorkType.PRACTICE,
        work_number=1,
        work_type_custom=None,
        discipline="Информатика",
        topic="Тестовый отчёт",
        student_full_name="Иванов Иван Иванович",
        group="ББИ-24-3",
        semester="2",
        direction_code="38.03.05",
        direction_name="Бизнес-информатика",
        department="Кафедра информатики",
        teacher_full_name="Петров Пётр Петрович",
        submission_date=date(2025, 5, 20),
    )

    main_section = SectionBlock(title="1 Основной раздел")
    main_section.children.append(
        ListBlock(list_type="numbered", items=["Первый пункт", "Второй пункт"])
    )
    main_section.children.append(
        TableBlock(
            caption="Таблица 1 – Пример данных",
            rows=[["Колонка 1", "Колонка 2"], ["1", "2"]],
        )
    )

    return Report(meta=meta, blocks=[main_section])


def test_build_docx_with_minimal_report_returns_valid_docx():
    report = _build_minimal_valid_report()

    data = build_docx(report)
    assert isinstance(data, (bytes, bytearray))
    assert len(data) > 0

    doc = Document(BytesIO(data))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Первая практическая работа" in full_text
    assert "Текст введения." in full_text


def test_build_docx_replaces_title_placeholders_from_template():
    template_doc = Document()
    template_doc.add_paragraph("Студент: {{STUDENT_NAME}}")
    template_doc.add_paragraph("Группа: {{GROUP}}")

    buffer = BytesIO()
    template_doc.save(buffer)
    template_bytes = buffer.getvalue()

    meta = ReportMeta(
        preset="misis_v1",
        work_type=WorkType.PRACTICE,
        work_number=1,
        work_type_custom=None,
        discipline="Информатика",
        topic="Первая практическая работа",
        student_full_name="Сидоров Сидор Сидорович",
        group="ББИ-24-3",
        semester="2",
        direction_code="38.03.05",
        direction_name="Бизнес-информатика",
        department="Кафедра информатики",
        teacher_full_name="Петров Пётр Петрович",
        submission_date=date(2025, 5, 20),
    )

    report = Report(meta=meta, blocks=[])

    output_bytes = build_docx(
        report,
        preset_id="misis_v1",
        title_template=template_bytes,
    )
    doc = Document(BytesIO(output_bytes))

    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "{{STUDENT_NAME}}" not in texts
    assert "{{GROUP}}" not in texts
    assert "Сидоров Сидор Сидорович" in texts
    assert "ББИ-24-3" in texts


def test_nested_blocks_are_rendered_in_docx():
    report = _build_sample_report_with_nested_blocks()

    data = build_docx(report)
    doc = Document(BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)

    assert "Краткое введение." in texts
    assert "Комментарий к рисунку и таблице." in texts
    assert "Текст приложения." in texts


def test_normal_and_headings_use_times_new_roman_and_black_color():
    report = _build_minimal_valid_report()
    data = build_docx(report)
    doc = Document(BytesIO(data))

    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size.pt == 12
    assert normal.font.color.rgb is not None
    assert normal.font.color.rgb == RGBColor(0, 0, 0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        assert style.font.name == "Times New Roman"
        assert style.font.size.pt == 12
        assert style.font.bold is True
        assert style.font.color.rgb is not None
        assert style.font.color.rgb == RGBColor(0, 0, 0)


def test_lists_and_table_are_rendered_with_numbers_and_captions():
    report = _build_sample_report_with_list_and_table()
    data = build_docx(report)
    doc = Document(BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)

    assert "1) Первый пункт" in texts
    assert "2) Второй пункт" in texts
    assert "Таблица 1 – Пример данных" in texts


def test_headings_have_numbering_properties():
    report = _build_minimal_valid_report()
    data = build_docx(report)
    doc = Document(BytesIO(data))

    style = doc.styles["Heading 1"]
    style_elm = style.element
    ppr = style_elm.pPr
    assert ppr is not None
    num_pr = ppr.numPr
    assert num_pr is not None
    assert num_pr.numId is not None
    assert num_pr.ilvl is not None
