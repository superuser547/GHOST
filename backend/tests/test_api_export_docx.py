from __future__ import annotations

import base64
from datetime import date
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.models import Report, ReportMeta, SectionBlock, TextBlock, WorkType

client = TestClient(app)


def _build_minimal_valid_report() -> Report:
    meta = ReportMeta(
        preset="misis_v1",
        work_type=WorkType.PRACTICE,
        work_number=1,
        work_type_custom=None,
        discipline="Информатика",
        topic="Первая практическая работа",
        student_full_name="Иванов Иван Иванович",
        group="ББИ-24-3",
        semester="2",
        direction_code="38.03.05",
        direction_name="Бизнес-информатика",
        department="Кафедра информатики",
        teacher_full_name="Петров Пётр Петрович",
        submission_date=date(2025, 5, 20),
    )

    intro = SectionBlock(title="Введение", special_kind="INTRO")
    main_section = SectionBlock(title="Основной раздел")
    conclusion = SectionBlock(title="Заключение", special_kind="CONCLUSION")

    blocks = [
        intro,
        TextBlock(text="Текст введения."),
        main_section,
        TextBlock(text="Основной текст."),
        conclusion,
    ]

    return Report(meta=meta, blocks=blocks)


def test_export_docx_success_for_valid_report():
    report = _build_minimal_valid_report()

    payload = {
        "report": report.model_dump(mode="json"),
        "preset_id": "misis_v1",
        "title_template": None,
    }

    response = client.post("/api/v1/reports/export/docx", json=payload)
    assert response.status_code == 200

    content_type = response.headers.get("content-type")
    assert (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    content_disposition = response.headers.get("content-disposition", "")
    assert "attachment;" in content_disposition
    assert ".docx" in content_disposition

    data = response.content
    assert isinstance(data, (bytes, bytearray))
    assert len(data) > 0

    doc = Document(BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Первая практическая работа" in texts
    assert "Текст введения." in texts


def test_export_docx_returns_422_when_validation_fails():
    report = _build_minimal_valid_report()
    report.blocks = []

    payload = {
        "report": report.model_dump(mode="json"),
        "preset_id": "misis_v1",
        "title_template": None,
    }

    response = client.post("/api/v1/reports/export/docx", json=payload)
    assert response.status_code == 422

    body = response.json()
    assert "detail" in body

    detail = body["detail"]
    assert isinstance(detail, dict)
    assert "errors" in detail
    assert isinstance(detail["errors"], list)
    assert detail["errors"], "Ожидаем хотя бы одну ошибку валидации"
    assert "warnings" in detail


def test_export_docx_uses_title_template_and_replaces_placeholders():
    template_doc = Document()
    template_doc.add_paragraph("Студент: {{STUDENT_NAME}}")
    template_doc.add_paragraph("Группа: {{GROUP}}")

    buffer = BytesIO()
    template_doc.save(buffer)
    template_bytes = buffer.getvalue()

    report = _build_minimal_valid_report()

    payload = {
        "report": report.model_dump(mode="json"),
        "preset_id": "misis_v1",
        "title_template": base64.b64encode(template_bytes).decode(),
    }

    response = client.post("/api/v1/reports/export/docx", json=payload)
    assert response.status_code == 200

    doc = Document(BytesIO(response.content))
    texts = "\n".join(p.text for p in doc.paragraphs)

    assert "{{STUDENT_NAME}}" not in texts
    assert "{{GROUP}}" not in texts
    assert "Иванов Иван Иванович" in texts
    assert "ББИ-24-3" in texts
